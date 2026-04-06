#!/usr/bin/env python3.11
"""
标书工厂 - 自动生成标书文档
根据招标公告信息自动生成技术标、商务标和资质文件清单
"""

import json
import re
import os
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn


class BidGenerator:
    """标书生成器"""
    
    def __init__(self, output_dir: str = None):
        self.output_dir = output_dir or os.path.join(os.path.dirname(__file__), '../outputs')
        self.templates_dir = os.path.join(os.path.dirname(__file__), '../templates')
        os.makedirs(self.output_dir, exist_ok=True)
        
        # 公司信息（需要用户配置）
        self.company_info = {
            'name': 'XX 科技有限公司',
            'address': 'XX 省 XX 市 XX 区 XX 路 XX 号',
            'phone': '010-XXXXXXXX',
            'email': 'bid@company.com',
            'legal_rep': 'XXX',
            'register_capital': 'XXXX 万元',
            'establish_date': '20XX 年 XX 月',
        }
    
    def extract_bid_info(self, announcement_text: str) -> Dict:
        """从招标公告中提取关键信息"""
        info = {
            'project_name': '',
            'project_number': '',
            'tenderer': '',
            'bid_deadline': '',
            'open_bid_time': '',
            'project_location': '',
            'budget': '',
            'project_scope': '',
            'qualification_requirements': [],
            'technical_requirements': [],
            'contact_info': {},
        }
        
        # 提取项目名称
        name_patterns = [
            r'项目名称.\s*(.+?)(?:\n|$)',
            r'一、项目名称.\s*(.+?)(?:\n|$)',
            r'工程名称.\s*(.+?)(?:\n|$)',
        ]
        for pattern in name_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['project_name'] = match.group(1).strip()
                break
        
        # 提取项目编号
        number_patterns = [
            r'项目编号.\s*(.+?)(?:\n|$)',
            r'招标编号.\s*(.+?)(?:\n|$)',
        ]
        for pattern in number_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['project_number'] = match.group(1).strip()
                break
        
        # 提取招标人
        tenderer_patterns = [
            r'招标人.\s*(.+?)(?:\n|$)',
            r'采购人.\s*(.+?)(?:\n|$)',
            r'业主单位.\s*(.+?)(?:\n|$)',
        ]
        for pattern in tenderer_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['tenderer'] = match.group(1).strip()
                break
        
        # 提取投标截止时间
        deadline_patterns = [
            r'投标截止时间.\s*(.+?)(?:\n|$)',
            r'递交投标文件截止时间.\s*(.+?)(?:\n|$)',
            r'截止时间.\s*(.+?)(?:\n|$)',
        ]
        for pattern in deadline_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['bid_deadline'] = match.group(1).strip()
                break
        
        # 提取预算金额
        budget_patterns = [
            r'预算金额.\s*([¥￥]?\s*\d+(?:,\d{3})*(?:\.\d+)?[万元]?)',
            r'招标控制价.\s*([¥￥]?\s*\d+(?:,\d{3})*(?:\.\d+)?[万元]?)',
            r'最高限价.\s*([¥￥]?\s*\d+(?:,\d{3})*(?:\.\d+)?[万元]?)',
        ]
        for pattern in budget_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['budget'] = match.group(1).strip()
                break
        
        # 提取项目地点
        location_patterns = [
            r'项目地点.\s*(.+?)(?:\n|$)',
            r'实施地点.\s*(.+?)(?:\n|$)',
            r'交货地点.\s*(.+?)(?:\n|$)',
        ]
        for pattern in location_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['project_location'] = match.group(1).strip()
                break
        
        # 提取联系方式
        contact_patterns = [
            r'联系人.\s*(.+?)(?:\n|$)',
            r'联系电话.\s*(.+?)(?:\n|$)',
            r'联系方式.\s*(.+?)(?:\n|$)',
        ]
        for pattern in contact_patterns:
            match = re.search(pattern, announcement_text)
            if match:
                info['contact_info']['name'] = match.group(1).strip()
                break
        
        return info
    
    def generate_technical_bid(self, bid_info: Dict, output_filename: str = None) -> str:
        """生成技术标文档"""
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
        
        # 封面
        self._add_cover_page(doc, bid_info, '技术标')
        
        # 目录
        doc.add_page_break()
        self._add_table_of_contents(doc)
        
        # 第一章 项目理解
        doc.add_page_break()
        self._add_chapter1(doc, bid_info)
        
        # 第二章 技术方案
        doc.add_page_break()
        self._add_chapter2(doc, bid_info)
        
        # 第三章 实施计划
        doc.add_page_break()
        self._add_chapter3(doc, bid_info)
        
        # 第四章 质量保证
        doc.add_page_break()
        self._add_chapter4(doc, bid_info)
        
        # 第五章 售后服务
        doc.add_page_break()
        self._add_chapter5(doc, bid_info)
        
        # 保存文档
        if output_filename is None:
            project_name = bid_info.get('project_name', '项目')[:20]
            output_filename = f"{project_name}_技术标_{datetime.now().strftime('%Y%m%d')}.docx"
        
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_business_bid(self, bid_info: Dict, output_filename: str = None) -> str:
        """生成商务标文档"""
        doc = Document()
        
        # 设置页面
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        
        # 封面
        self._add_cover_page(doc, bid_info, '商务标')
        
        # 目录
        doc.add_page_break()
        self._add_business_toc(doc)
        
        # 第一章 投标函
        doc.add_page_break()
        self._add_business_chapter1(doc, bid_info)
        
        # 第二章 报价表
        doc.add_page_break()
        self._add_business_chapter2(doc, bid_info)
        
        # 第三章 商务条款响应
        doc.add_page_break()
        self._add_business_chapter3(doc, bid_info)
        
        # 第四章 企业资质
        doc.add_page_break()
        self._add_business_chapter4(doc, bid_info)
        
        # 第五章 业绩案例
        doc.add_page_break()
        self._add_business_chapter5(doc, bid_info)
        
        # 保存文档
        if output_filename is None:
            project_name = bid_info.get('project_name', '项目')[:20]
            output_filename = f"{project_name}_商务标_{datetime.now().strftime('%Y%m%d')}.docx"
        
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_qualification_list(self, bid_info: Dict, output_filename: str = None) -> str:
        """生成资质文件清单"""
        doc = Document()
        
        # 标题
        title = doc.add_heading('资质文件清单', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 项目信息
        doc.add_paragraph(f'项目名称：{bid_info.get("project_name", "")}')
        doc.add_paragraph(f'项目编号：{bid_info.get("project_number", "")}')
        doc.add_paragraph(f'生成日期：{datetime.now().strftime("%Y 年%m 月%d 日")}')
        
        doc.add_paragraph()
        
        # 资质文件清单表格
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        header_cells = table.rows[0].cells
        headers = ['序号', '文件名称', '要求', '状态']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
        
        # 必备资质
        qualifications = [
            ('1', '营业执照副本复印件', '必须提供，加盖公章', '□'),
            ('2', '法定代表人身份证明', '必须提供', '□'),
            ('3', '法定代表人授权委托书', '如有授权，必须提供', '□'),
            ('4', '税务登记证复印件', '必须提供，加盖公章', '□'),
            ('5', '组织机构代码证复印件', '必须提供，加盖公章', '□'),
            ('6', '开户许可证复印件', '必须提供，加盖公章', '□'),
        ]
        
        for qual in qualifications:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                cell.text = qual[i]
        
        doc.add_paragraph()
        doc.add_heading('专业资质（根据项目类型提供）', level=2)
        
        professional_table = doc.add_table(rows=1, cols=4)
        professional_table.style = 'Table Grid'
        
        prof_headers = ['序号', '文件名称', '要求', '状态']
        for i, header in enumerate(prof_headers):
            professional_table.rows[0].cells[i].text = header
            professional_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        professional_quas = [
            ('1', '相关行业资质证书', '根据招标要求', '□'),
            ('2', 'ISO 质量管理体系认证', '优先提供', '□'),
            ('3', '安全生产许可证', '工程类项目必须', '□'),
            ('4', '专业人员资格证书', '根据项目需要', '□'),
            ('5', '类似项目业绩证明', '建议提供近 3 年', '□'),
        ]
        
        for qual in professional_quas:
            row = professional_table.add_row()
            for i, cell in enumerate(row.cells):
                cell.text = qual[i]
        
        # 保存文档
        if output_filename is None:
            project_name = bid_info.get('project_name', '项目')[:20]
            output_filename = f"{project_name}_资质文件清单_{datetime.now().strftime('%Y%m%d')}.docx"
        
        output_path = os.path.join(self.output_dir, output_filename)
        doc.save(output_path)
        
        return output_path
    
    def generate_all(self, announcement_text: str, project_name: str = None) -> Dict[str, str]:
        """生成完整标书（技术标 + 商务标 + 资质清单）"""
        bid_info = self.extract_bid_info(announcement_text)
        
        if project_name:
            bid_info['project_name'] = project_name
        
        results = {}
        
        # 生成技术标
        results['technical'] = self.generate_technical_bid(bid_info)
        
        # 生成商务标
        results['business'] = self.generate_business_bid(bid_info)
        
        # 生成资质清单
        results['qualification'] = self.generate_qualification_list(bid_info)
        
        return results
    
    def _add_cover_page(self, doc: Document, bid_info: Dict, bid_type: str):
        """添加封面页"""
        # 项目名称
        project_name = bid_info.get('project_name', '项目名称')
        title = doc.add_heading(project_name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 标书类型
        type_para = doc.add_paragraph(f'\n\n{bid_type}\n\n')
        type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        type_para.runs[0].font.size = Pt(36)
        type_para.runs[0].font.bold = True
        
        # 投标人信息
        doc.add_paragraph('\n\n\n')
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f'投标人：{self.company_info["name"]}\n')
        info_para.add_run(f'日期：{datetime.now().strftime("%Y 年%m 月%d 日")}\n')
    
    def _add_table_of_contents(self, doc: Document):
        """添加技术标目录"""
        doc.add_heading('目录', level=1)
        
        chapters = [
            '第一章 项目理解',
            '  1.1 项目背景',
            '  1.2 项目需求分析',
            '  1.3 项目目标',
            '第二章 技术方案',
            '  2.1 技术路线',
            '  2.2 系统设计',
            '  2.3 功能模块',
            '第三章 实施计划',
            '  3.1 项目进度安排',
            '  3.2 人员配置',
            '  3.3 资源配置',
            '第四章 质量保证',
            '  4.1 质量管理体系',
            '  4.2 质量控制措施',
            '  4.3 验收标准',
            '第五章 售后服务',
            '  5.1 服务承诺',
            '  5.2 培训计划',
            '  5.3 维护方案',
        ]
        
        for chapter in chapters:
            doc.add_paragraph(chapter)
    
    def _add_chapter1(self, doc: Document, bid_info: Dict):
        """第一章：项目理解"""
        doc.add_heading('第一章 项目理解', level=1)
        
        doc.add_heading('1.1 项目背景', level=2)
        doc.add_paragraph(f'本项目为{bid_info.get("project_name", "该项目")}，招标人为{bid_info.get("tenderer", "招标人")}。')
        doc.add_paragraph('我方充分理解本项目的重要性和紧迫性，将全力以赴确保项目成功实施。')
        
        doc.add_heading('1.2 项目需求分析', level=2)
        doc.add_paragraph('根据招标文件要求，本项目主要需求包括：')
        doc.add_paragraph('1. 满足招标文件规定的技术要求')
        doc.add_paragraph('2. 符合相关行业标准和规范')
        doc.add_paragraph('3. 保证项目质量和进度')
        
        doc.add_heading('1.3 项目目标', level=2)
        doc.add_paragraph('本项目目标：')
        doc.add_paragraph('• 按时保质完成项目实施')
        doc.add_paragraph('• 满足招标人全部技术要求')
        doc.add_paragraph('• 提供优质的售后服务')
    
    def _add_chapter2(self, doc: Document, bid_info: Dict):
        """第二章：技术方案"""
        doc.add_heading('第二章 技术方案', level=1)
        
        doc.add_heading('2.1 技术路线', level=2)
        doc.add_paragraph('我方将采用成熟、先进的技术路线，确保系统稳定可靠。')
        doc.add_paragraph('技术选型原则：先进性、可靠性、经济性、可扩展性。')
        
        doc.add_heading('2.2 系统设计', level=2)
        doc.add_paragraph('系统架构设计遵循模块化、分层化原则，便于维护和扩展。')
        
        doc.add_heading('2.3 功能模块', level=2)
        doc.add_paragraph('根据项目需求，系统将包含以下功能模块：')
        doc.add_paragraph('1. 核心业务模块')
        doc.add_paragraph('2. 数据管理模块')
        doc.add_paragraph('3. 用户管理模块')
        doc.add_paragraph('4. 系统管理模块')
    
    def _add_chapter3(self, doc: Document, bid_info: Dict):
        """第三章：实施计划"""
        doc.add_heading('第三章 实施计划', level=1)
        
        doc.add_heading('3.1 项目进度安排', level=2)
        doc.add_paragraph('项目总工期：根据招标文件要求')
        doc.add_paragraph('关键节点：')
        doc.add_paragraph('• 合同签订后 X 日内完成需求调研')
        doc.add_paragraph('• 需求确认后 X 日内完成系统设计')
        doc.add_paragraph('• 设计确认后 X 日内完成开发实施')
        doc.add_paragraph('• 开发完成后 X 日内完成测试验收')
        
        doc.add_heading('3.2 人员配置', level=2)
        doc.add_paragraph('项目团队配置：')
        doc.add_paragraph('• 项目经理：1 名（负责整体协调）')
        doc.add_paragraph('• 技术负责人：1 名（负责技术方案）')
        doc.add_paragraph('• 开发工程师：若干名（负责开发实施）')
        doc.add_paragraph('• 测试工程师：若干名（负责质量测试）')
        
        doc.add_heading('3.3 资源配置', level=2)
        doc.add_paragraph('项目所需资源将按计划及时到位，确保项目顺利实施。')
    
    def _add_chapter4(self, doc: Document, bid_info: Dict):
        """第四章：质量保证"""
        doc.add_heading('第四章 质量保证', level=1)
        
        doc.add_heading('4.1 质量管理体系', level=2)
        doc.add_paragraph('我方已通过 ISO9001 质量管理体系认证，建立了完善的质量管理体系。')
        
        doc.add_heading('4.2 质量控制措施', level=2)
        doc.add_paragraph('质量控制措施包括：')
        doc.add_paragraph('• 需求评审：确保需求理解准确')
        doc.add_paragraph('• 设计评审：确保技术方案合理')
        doc.add_paragraph('• 代码审查：确保代码质量')
        doc.add_paragraph('• 测试验证：确保功能完整')
        
        doc.add_heading('4.3 验收标准', level=2)
        doc.add_paragraph('项目验收标准严格按照招标文件和合同约定执行。')
    
    def _add_chapter5(self, doc: Document, bid_info: Dict):
        """第五章：售后服务"""
        doc.add_heading('第五章 售后服务', level=1)
        
        doc.add_heading('5.1 服务承诺', level=2)
        doc.add_paragraph('我方承诺：')
        doc.add_paragraph('• 质保期：X 年（自验收合格之日起）')
        doc.add_paragraph('• 响应时间：7×24 小时响应，X 小时内到达现场')
        doc.add_paragraph('• 问题解决：一般问题 X 小时内解决，重大问题 X 小时内解决')
        
        doc.add_heading('5.2 培训计划', level=2)
        doc.add_paragraph('提供全面的培训服务：')
        doc.add_paragraph('• 系统操作培训')
        doc.add_paragraph('• 系统维护培训')
        doc.add_paragraph('• 管理员培训')
        
        doc.add_heading('5.3 维护方案', level=2)
        doc.add_paragraph('维护服务包括：')
        doc.add_paragraph('• 定期巡检')
        doc.add_paragraph('• 系统升级')
        doc.add_paragraph('• 故障处理')
        doc.add_paragraph('• 技术咨询')
    
    def _add_business_toc(self, doc: Document):
        """添加商务标目录"""
        doc.add_heading('目录', level=1)
        
        chapters = [
            '第一章 投标函',
            '第二章 报价表',
            '第三章 商务条款响应',
            '第四章 企业资质',
            '第五章 业绩案例',
        ]
        
        for chapter in chapters:
            doc.add_paragraph(chapter)
    
    def _add_business_chapter1(self, doc: Document, bid_info: Dict):
        """商务标第一章：投标函"""
        doc.add_heading('第一章 投标函', level=1)
        
        content = f"""致：{bid_info.get('tenderer', '招标人')}

我方已仔细阅读了{bid_info.get('project_name', '本项目')}的招标文件，决定参加投标。

一、我方愿意按照招标文件规定的各项要求，向招标人提供所需的货物/服务。

二、我方承诺：
1. 严格遵守招标文件的全部条款
2. 保证所提供的货物/服务符合质量要求
3. 按时履行合同义务

三、本投标函有效期为投标截止日后 XX 天。

投标人：{self.company_info['name']}（盖章）

法定代表人或授权代表：（签字）

日期：{datetime.now().strftime('%Y 年%m 月%d 日')}
"""
        doc.add_paragraph(content)
    
    def _add_business_chapter2(self, doc: Document, bid_info: Dict):
        """商务标第二章：报价表"""
        doc.add_heading('第二章 报价表', level=1)
        
        doc.add_paragraph(f'项目名称：{bid_info.get("project_name", "")}')
        doc.add_paragraph(f'项目编号：{bid_info.get("project_number", "")}')
        doc.add_paragraph()
        
        # 报价表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['序号', '项目名称', '规格型号', '数量', '报价（元）']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        # 示例数据行
        items = [
            ('1', '项目服务', '按招标要求', '1 项', '待填写'),
            ('2', '设备材料', '按招标要求', '1 批', '待填写'),
            ('3', '实施费用', '按招标要求', '1 项', '待填写'),
            ('4', '培训费用', '按招标要求', '1 项', '待填写'),
            ('5', '售后服务', '按招标要求', '1 项', '待填写'),
        ]
        
        for item in items:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                cell.text = item[i]
        
        # 合计
        total_row = table.add_row()
        total_row.cells[0].text = ''
        total_row.cells[1].text = '合计'
        total_row.cells[2].text = ''
        total_row.cells[3].text = ''
        total_row.cells[4].text = '待填写'
        
        doc.add_paragraph()
        doc.add_paragraph('注：以上报价为含税价，包含完成本项目所需的全部费用。')
    
    def _add_business_chapter3(self, doc: Document, bid_info: Dict):
        """商务标第三章：商务条款响应"""
        doc.add_heading('第三章 商务条款响应', level=1)
        
        doc.add_heading('3.1 付款方式响应', level=2)
        doc.add_paragraph('我方完全响应招标文件规定的付款方式。')
        
        doc.add_heading('3.2 交货/实施周期响应', level=2)
        doc.add_paragraph('我方承诺按招标文件要求的时间完成交货/实施。')
        
        doc.add_heading('3.3 质保期响应', level=2)
        doc.add_paragraph('我方提供的质保期不低于招标文件要求。')
        
        doc.add_heading('3.4 其他商务条款', level=2)
        doc.add_paragraph('我方对招标文件中的其他商务条款均无异议，完全响应。')
    
    def _add_business_chapter4(self, doc: Document, bid_info: Dict):
        """商务标第四章：企业资质"""
        doc.add_heading('第四章 企业资质', level=1)
        
        doc.add_heading('4.1 企业基本信息', level=2)
        doc.add_paragraph(f'企业名称：{self.company_info["name"]}')
        doc.add_paragraph(f'注册地址：{self.company_info["address"]}')
        doc.add_paragraph(f'法定代表人：{self.company_info["legal_rep"]}')
        doc.add_paragraph(f'注册资本：{self.company_info["register_capital"]}')
        doc.add_paragraph(f'成立日期：{self.company_info["establish_date"]}')
        
        doc.add_heading('4.2 资质证书', level=2)
        doc.add_paragraph('以下资质证书复印件附后：')
        doc.add_paragraph('• 营业执照')
        doc.add_paragraph('• 税务登记证')
        doc.add_paragraph('• 组织机构代码证')
        doc.add_paragraph('• 相关行业资质证书')
        
        doc.add_heading('4.3 财务状况', level=2)
        doc.add_paragraph('近三年财务审计报告附后。')
    
    def _add_business_chapter5(self, doc: Document, bid_info: Dict):
        """商务标第五章：业绩案例"""
        doc.add_heading('第五章 业绩案例', level=1)
        
        doc.add_paragraph('近三年类似项目业绩：')
        
        # 业绩表格
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        
        headers = ['序号', '项目名称', '业主单位', '合同金额', '完成时间']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
        # 示例业绩
        projects = [
            ('1', 'XXX 项目', 'XXX 单位', 'XXX 万元', '202X 年'),
            ('2', 'XXX 项目', 'XXX 单位', 'XXX 万元', '202X 年'),
            ('3', 'XXX 项目', 'XXX 单位', 'XXX 万元', '202X 年'),
        ]
        
        for project in projects:
            row = table.add_row()
            for i, cell in enumerate(row.cells):
                cell.text = project[i]
        
        doc.add_paragraph()
        doc.add_paragraph('注：业绩证明材料（合同复印件、验收报告等）附后。')


def main():
    """主函数 - 示例用法"""
    generator = BidGenerator()
    
    # 示例招标公告文本
    sample_announcement = """
    项目名称：某市电力设备采购项目
    项目编号：SGCC-2024-001
    招标人：国网某省电力公司
    预算金额：500 万元
    投标截止时间：2024 年 X 月 X 日
    项目地点：某省某市
    联系人：张老师
    联系电话：010-XXXXXXXX
    """
    
    # 生成标书
    results = generator.generate_all(sample_announcement)
    
    print("标书生成完成：")
    for bid_type, path in results.items():
        print(f"  {bid_type}: {path}")


if __name__ == '__main__':
    main()
